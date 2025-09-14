from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Inches
from io import BytesIO
from flask_cors import CORS

app = Flask(__name__)
CORS(app) # Autoriser les requêtes cross-origin

@app.route("/api/generate-word", methods=["POST"])
def generate_word():
    data = request.json

    document = Document()

    document.add_heading("Cahier des Charges - " + data["projectName"], 0)

    document.add_heading("1. Présentation du Projet", level=1)
    document.add_paragraph("Nom du projet : {}".format(data["projectName"]))
    document.add_paragraph("Objectif principal : {}".format(data["mainObjective"]))
    document.add_paragraph("Type de produit : {}".format(data["productType"]))
    document.add_paragraph("Vision à long terme : {}".format(data["longTermVision"]))

    document.add_heading("2. Analyse du Marché et Utilisateurs", level=1)
    document.add_paragraph("Public cible : {}".format(data["targetAudience"]))
    document.add_paragraph("Problèmes utilisateurs à résoudre : {}".format(data["userProblems"]))
    document.add_paragraph("Proposition de valeur unique : {}".format(data["uniqueValueProposition"]))
    document.add_paragraph("Objectifs SMART :")
    for obj in data["smartObjectives"]:
        document.add_paragraph("- {}".format(obj), style="List Bullet")
    document.add_paragraph("Parties prenantes :")
    for stakeholder in data["stakeholders"]:
        document.add_paragraph("- {}".format(stakeholder), style="List Bullet")

    document.add_heading("3. Spécifications Fonctionnelles", level=1)
    document.add_paragraph("Fonctionnalités MVP (essentielles) :")
    for feature in data["mvpFeatures"]:
        document.add_paragraph("- {}".format(feature), style="List Bullet")
    document.add_paragraph("Fonctionnalités futures :")
    for feature in data["futureFeatures"]:
        document.add_paragraph("- {}".format(feature), style="List Bullet")
    document.add_paragraph("Intégrations nécessaires :")
    for integration in data["integrations"]:
        document.add_paragraph("- {}".format(integration), style="List Bullet")

    document.add_heading("4. Spécifications Techniques", level=1)
    document.add_paragraph("Exigences UX/UI : {}".format(data["uxUiRequirements"]))
    document.add_paragraph("Wireframes/Maquettes nécessaires : {}".format("Oui" if data["wireframesNeeded"] else "Non"))
    document.add_paragraph("Charte graphique/Marque : {}".format(data["graphicCharter"]))
    document.add_paragraph("Exigences de performance : {}".format(data["performanceRequirements"]))
    document.add_paragraph("Exigences de sécurité : {}".format(data["securityRequirements"]))
    document.add_paragraph("Compatibilité : {}".format(data["compatibility"]))

    document.add_heading("5. Gestion de Projet", level=1)
    document.add_paragraph("Méthodologie de projet : {}".format(data["projectMethodology"]))
    if data["projectMethodology"] == "Agile/Scrum":
        document.add_paragraph("Fréquence des sprints : {}".format(data["sprintFrequency"]))
    document.add_paragraph("Ressources humaines nécessaires :")
    for resource in data["humanResources"]:
        document.add_paragraph("- {}".format(resource), style="List Bullet")
    document.add_paragraph("Budget estimé : {}".format(data["estimatedBudget"]))
    document.add_paragraph("Délai souhaité : {}".format(data["desiredTimeline"]))
    document.add_paragraph("Outils de suivi préférés : {}".format(data["preferredTrackingTools"]))
    document.add_paragraph("Risques identifiés :")
    for risk in data["identifiedRisks"]:
        document.add_paragraph("- {}".format(risk), style="List Bullet")
    document.add_paragraph("Stratégies d'atténuation : {}".format(data["mitigationStrategies"]))

    document.add_heading("6. Lancement et Suivi", level=1)
    document.add_paragraph("Plan de lancement : {}".format(data["launchPlan"]))
    document.add_paragraph("Date de lancement souhaitée : {}".format(data["desiredLaunchDate"]))
    document.add_paragraph("Phase de test bêta : {}".format("Oui" if data["betaTesting"] else "Non"))
    document.add_paragraph("Canaux de support :")
    for channel in data["supportChannels"]:
        document.add_paragraph("- {}".format(channel), style="List Bullet")
    document.add_paragraph("Plan de maintenance : {}".format(data["maintenancePlan"]))
    document.add_paragraph("Processus d'amélioration continue : {}".format(data["continuousImprovementProcess"]))
    document.add_paragraph("KPIs à suivre :")
    for kpi in data["successMetrics"]:
        document.add_paragraph("- {}".format(kpi), style="List Bullet")

    document.add_heading("Signatures", level=1)
    document.add_paragraph("\n\n")
    document.add_paragraph("_________________________\nSignature Client")
    document.add_paragraph("\n\n")
    document.add_paragraph("_________________________\nSignature Entreprise")

    # Sauvegarder le document en mémoire
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name="cahier_des_charges.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/api/generate-excel", methods=["POST"])
def generate_excel():
    data = request.json
    # Ici, vous implémenteriez la logique de génération Excel avec diagramme de Gantt
    # Pour l'instant, nous renvoyons un fichier CSV simple
    from generate_excel_gantt import generate_excel_gantt
    file_stream = generate_excel_gantt(data)
    return send_file(file_stream, as_attachment=True, download_name="plan_action_gantt.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)


