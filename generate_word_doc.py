
from docx import Document
from docx.shared import Inches

def generate_word_doc(data):
    document = Document()

    document.add_heading("Cahier des Charges - " + data["projectName"], 0)

    document.add_heading("1. Présentation du Projet", level=1)
    document.add_paragraph(f"Nom du projet : {data["projectName"]}")
    document.add_paragraph(f"Objectif principal : {data["mainObjective"]}")
    document.add_paragraph(f"Type de produit : {data["productType"]}")
    document.add_paragraph(f"Vision à long terme : {data["longTermVision"]}")

    document.add_heading("2. Analyse du Marché et Utilisateurs", level=1)
    document.add_paragraph(f"Public cible : {data["targetAudience"]}")
    document.add_paragraph(f"Problèmes utilisateurs à résoudre : {data["userProblems"]}")
    document.add_paragraph(f"Proposition de valeur unique : {data["uniqueValueProposition"]}")
    document.add_paragraph("Objectifs SMART :")
    for obj in data["smartObjectives"]:
        document.add_paragraph(f"- {obj}", style=\


t("List Bullet"))
    document.add_paragraph("Parties prenantes :")
    for stakeholder in data["stakeholders"]:
        document.add_paragraph(f"- {stakeholder}", style="List Bullet")

    document.add_heading("3. Spécifications Fonctionnelles", level=1)
    document.add_paragraph("Fonctionnalités MVP (essentielles) :")
    for feature in data["mvpFeatures"]:
        document.add_paragraph(f"- {feature}", style="List Bullet")
    document.add_paragraph("Fonctionnalités futures :")
    for feature in data["futureFeatures"]:
        document.add_paragraph(f"- {feature}", style="List Bullet")
    document.add_paragraph("Intégrations nécessaires :")
    for integration in data["integrations"]:
        document.add_paragraph(f"- {integration}", style="List Bullet")

    document.add_heading("4. Spécifications Techniques", level=1)
    document.add_paragraph(f"Exigences UX/UI : {data["uxUiRequirements"]}")
    document.add_paragraph(f"Wireframes/Maquettes nécessaires : {"Oui" if data["wireframesNeeded"] else "Non"}")
    document.add_paragraph(f"Charte graphique/Marque : {data["graphicCharter"]}")
    document.add_paragraph(f"Exigences de performance : {data["performanceRequirements"]}")
    document.add_paragraph(f"Exigences de sécurité : {data["securityRequirements"]}")
    document.add_paragraph(f"Compatibilité : {data["compatibility"]}")

    document.add_heading("5. Gestion de Projet", level=1)
    document.add_paragraph(f"Méthodologie de projet : {data["projectMethodology"]}")
    if data["projectMethodology"] == "Agile/Scrum":
        document.add_paragraph(f"Fréquence des sprints : {data["sprintFrequency"]}")
    document.add_paragraph("Ressources humaines nécessaires :")
    for resource in data["humanResources"]:
        document.add_paragraph(f"- {resource}", style="List Bullet")
    document.add_paragraph(f"Budget estimé : {data["estimatedBudget"]}")
    document.add_paragraph(f"Délai souhaité : {data["desiredTimeline"]}")
    document.add_paragraph(f"Outils de suivi préférés : {data["preferredTrackingTools"]}")
    document.add_paragraph("Risques identifiés :")
    for risk in data["identifiedRisks"]:
        document.add_paragraph(f"- {risk}", style="List Bullet")
    document.add_paragraph(f"Stratégies d\"atténuation : {data["mitigationStrategies"]}")

    document.add_heading("6. Lancement et Suivi", level=1)
    document.add_paragraph(f"Plan de lancement : {data["launchPlan"]}")
    document.add_paragraph(f"Date de lancement souhaitée : {data["desiredLaunchDate"]}")
    document.add_paragraph(f"Phase de test bêta : {"Oui" if data["betaTesting"] else "Non"}")
    document.add_paragraph("Canaux de support :")
    for channel in data["supportChannels"]:
        document.add_paragraph(f"- {channel}", style="List Bullet")
    document.add_paragraph(f"Plan de maintenance : {data["maintenancePlan"]}")
    document.add_paragraph(f"Processus d\"amélioration continue : {data["continuousImprovementProcess"]}")
    document.add_paragraph("KPIs à suivre :")
    for kpi in data["successMetrics"]:
        document.add_paragraph(f"- {kpi}", style="List Bullet")

    document.add_heading("Signatures", level=1)
    document.add_paragraph("\n\n")
    document.add_paragraph("_________________________\nSignature Client")
    document.add_paragraph("\n\n")
    document.add_paragraph("_________________________\nSignature Entreprise")

    document.save("cahier_des_charges.docx")

if __name__ == "__main__":
    # Exemple de données (à remplacer par les données réelles de l'application React)
    example_data = {
        "projectName": "Plateforme E-Learning Entreprise",
        "mainObjective": "Créer une plateforme de formation en ligne moderne et intuitive pour les entreprises, permettant de former efficacement leurs employés avec des contenus personnalisés et un suivi des progrès en temps réel.",
        "longTermVision": "Devenir la référence en formation digitale en France",
        "targetAudience": "Professionnels de 25-45 ans cherchant à se former",
        "userProblems": "Manque de temps, formations trop chères, contenu pas adapté",
        "uniqueValueProposition": "Formations courtes, pratiques et certifiantes",
        "smartObjectives": ["Atteindre 1000 utilisateurs en 6 mois"],
        "stakeholders": ["Directeur Marketing - Validation du contenu"],
        "productType": "Application Web (SaaS)",
        "uxUiRequirements": "Interface intuitive, accessible, responsive",
        "wireframesNeeded": True,
        "graphicCharter": "Couleurs corporate, logo existant, style moderne",
        "performanceRequirements": "Temps de chargement < 3s, support 1000 utilisateurs simultanés",
        "securityRequirements": "Authentification 2FA, chiffrement des données, RGPD",
        "compatibility": "Chrome, Firefox, Safari, iOS, Android",
        "mvpFeatures": ["Inscription/Connexion utilisateur"],
        "futureFeatures": ["Système de recommandations IA"],
        "integrations": ["API de paiement Stripe, CRM Salesforce"],
        "projectMethodology": "Agile/Scrum",
        "sprintFrequency": "2 semaines",
        "humanResources": ["Chef de projet", "Développeurs front-end", "Développeurs back-end", "Designer UX/UI", "Testeur QA"],
        "estimatedBudget": "50 000 - 100 000 EUR",
        "desiredTimeline": "6 mois",
        "preferredTrackingTools": "Jira, Trello",
        "identifiedRisks": ["Retard de développement", "Manque d\"adhésion utilisateur"],
        "mitigationStrategies": "Communication régulière, tests utilisateurs fréquents",
        "launchPlan": "Lancement progressif par département",
        "desiredLaunchDate": "2025-12-01",
        "betaTesting": True,
        "supportChannels": ["Email", "Chat en ligne"],
        "maintenancePlan": "Mises à jour mensuelles",
        "continuousImprovementProcess": "Collecte de feedback continue",
        "successMetrics": ["Taux d\"engagement utilisateur", "Nombre de formations complétées"]
    }
    generate_word_doc(example_data)


